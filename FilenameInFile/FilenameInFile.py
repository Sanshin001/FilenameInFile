
import docx
import os
import re
import glob
import natsort
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def CreateDocx(name,string) :
    if name == "" :
        return
    file = docx.Document()
    for i in string :
        file.add_paragraph(i)

    file.save(name + ".docx")

def main() :
    Contents = []
    file_list = []
    title = input("파일 이름을 입력하세요 :")

    path = input("폴더 경로 입력(절대 경로 입력) : ")

    #폴더 내부의 파일 전체 읽기
    file_list = os.listdir(path);

    # 파일 1부터 정렬
    file_list = natsort.natsorted(file_list,reverse=False)
    #각 파일마다
    for file in file_list :
        #만약 docx 포맷이면
        if (".docx" in file) :
            #file 이름 추가하고
            Contents.append(file);
            #해당 파일을 열어서
            Originalfile = docx.Document(path+ "\\" + file)
            print(path+ "\\" + file)
            #한줄마다 추가한다.
            for x,paragraph in enumerate(Originalfile.paragraphs):
                Contents.append(paragraph.text)

    #만들어진 내용으로 docx 생성
    CreateDocx(title,Contents)

if __name__ == "__main__" :
    main()